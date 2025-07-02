# ------------- MedConsultant v8 (с авто-определением даты) -------------
import io, base64, datetime, json, os, sys, tempfile, uuid, subprocess, glob, re
import streamlit as st
import openai
import fitz  # PyMuPDF
import docx
from docx.shared import Pt, RGBColor
from docx2pdf import convert
from PIL import Image
import pydicom

# ---------- Инициализация API‑ключа ----------
# Ищем сначала в Streamlit Secrets, потом — в переменной окружения
openai.api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
if not openai.api_key:
    st.error("❗ Не задан OPENAI_API_KEY. Задайте секрет в Streamlit Cloud или переменную окружения.")
    st.stop()

# Создаём клиент OpenAI (будет использовать openai.api_key)
openai_client = openai.OpenAI()

# ---------- Конфигурация модели ----------
MODEL = "gpt-4o"
TEMPERATURE = 0.35
MAX_TOKENS = 4000
THEME = RGBColor(0, 102, 204)

# ---------- UI ----------
st.set_page_config("MedConsultant 🩺", page_icon="🩺", layout="wide")
st.title("🩺 MedConsultant — медицинский эксперт‑консультант")

st.sidebar.header("🧑‍⚕️ Данные пациента")
patient_name = st.sidebar.text_input("ФИО пациента")
patient_age = st.sidebar.text_input("Возраст")
patient_sex = st.sidebar.selectbox("Пол", ["", "М", "Ж"])

# ---------- Сессия ----------
if "files" not in st.session_state:
    st.session_state.files = {}


# ---------- НОВЫЙ ХЕЛПЕР: Авто-определение даты ----------
def extract_date_from_file(fname: str, data: bytes) -> datetime.date | None:
    """Пытается извлечь дату сначала из имени файла, затем из метаданных."""

    # 1. Поиск в имени файла (самый высокий приоритет)
    # Ищем форматы ГГГГ-ММ-ДД, ДД.ММ.ГГГГ, ДД-ММ-ГГГГ и т.д.
    patterns = [
        r'(\d{4})[._-](\d{2})[._-](\d{2})',  # 2025-07-02
        r'(\d{2})[._-](\d{2})[._-](\d{4})'  # 02-07-2025
    ]
    for pattern in patterns:
        match = re.search(pattern, fname)
        if match:
            try:
                parts = match.groups()
                if len(parts[0]) == 4:  # ГГГГ-ММ-ДД
                    year, month, day = map(int, parts)
                else:  # ДД-ММ-ГГГГ
                    day, month, year = map(int, parts)
                return datetime.date(year, month, day)
            except ValueError:
                continue

    # 2. Поиск в метаданных файла (если в имени не нашли)
    ext = fname.lower().split('.')[-1]
    try:
        if ext == "pdf":
            with fitz.open(stream=data, filetype="pdf") as doc:
                meta = doc.metadata
                date_str = meta.get('creationDate') or meta.get('modDate')
                if date_str:  # Формат D:YYYYMMDD...
                    return datetime.datetime.strptime(date_str[2:10], "%Y%m%d").date()
        elif ext == "docx":
            with io.BytesIO(data) as docx_io:
                doc = docx.Document(docx_io)
                return doc.core_properties.created.date()
        elif ext == "dcm":
            with io.BytesIO(data) as dcm_io:
                dcm = pydicom.dcmread(dcm_io)
                date_str = dcm.get("StudyDate") or dcm.get("AcquisitionDate")
                if date_str:
                    return datetime.datetime.strptime(date_str, "%Y%m%d").date()
    except Exception:
        return None  # Ошибка при чтении метаданных

    return None  # Если ничего не найдено


# ---------- Загрузка файлов ----------
uploaded = st.file_uploader(
    "Загрузите мед‑файлы (PDF, DOCX, изображения, DICOM…)",
    accept_multiple_files=True
)
if uploaded:
    for f in uploaded:
        if f.name not in st.session_state.files:
            file_data = f.read()
            # Пытаемся угадать дату, если не вышло - ставим сегодняшнюю
            guessed_date = extract_date_from_file(f.name, file_data) or datetime.date.today()
            st.session_state.files[f.name] = {
                "data": file_data,
                "note": "",
                "date": guessed_date  # Используем угаданную дату
            }

# ---------- Файлы, заметки и даты (с ручной корректировкой) ----------
if st.session_state.files:
    st.subheader("📁 Файлы, заметки и даты анализов")
    st.info("ℹ️ Дата для каждого файла определяется автоматически. Вы можете исправить её вручную.")

    # Сортируем файлы по дате для анализа в динамике
    sorted_files = sorted(st.session_state.files.items(), key=lambda item: item[1]['date'])
    st.session_state.files = dict(sorted_files)

    for fname, meta in st.session_state.files.items():
        with st.expander(f"{fname} (дата: {meta['date']:%d.%m.%Y})", expanded=False):
            col1, col2 = st.columns(2)
            with col1:
                meta["date"] = st.date_input(
                    "Дата анализа (можно исправить)",
                    value=meta["date"],
                    key=f"date_{fname}"
                )
            with col2:
                meta["note"] = st.text_area(
                    "Заметка к файлу (НЕ попадёт в отчёт)",
                    meta["note"],
                    key=f"note_{fname}",
                    placeholder="Например: это сцинтиграфия после нагрузки."
                )

st.markdown("---")
global_note = st.text_area(
    "📝 Общие указания для анализа (НЕ попадут в отчёт)",
    placeholder="Например: уделите внимание функции почек, сравните анализы крови за последние два месяца..."
)


# ---------- Основные функции (process_file, integrated_analysis, build_docx, docx_to_pdf) без изменений... ----------
# ... (вставьте сюда без изменений функции process_file, integrated_analysis, build_docx, docx_to_pdf из предыдущего ответа) ...
def process_file(fname: str, data: bytes) -> dict:
    """Извлекает текст и ИЗОБРАЖЕНИЯ из разных форматов файлов."""
    ext = fname.lower().split('.')[-1]
    text_content = ""
    images_base64 = []

    try:
        if ext == "pdf":
            doc = fitz.open(stream=data, filetype="pdf")
            for page in doc:
                text_content += page.get_text() + "\n"
                for img_instance in page.get_images(full=True):
                    xref = img_instance[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    images_base64.append(base64.b64encode(image_bytes).decode('utf-8'))
            doc.close()

        elif ext == "docx":
            doc = docx.Document(io.BytesIO(data))
            # Извлечение текста
            for para in doc.paragraphs:
                text_content += para.text + "\n"
            # Извлечение изображений
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_bytes = rel.target_part.blob
                    images_base64.append(base64.b64encode(image_bytes).decode('utf-8'))

        elif ext in ("txt", "csv", "md"):
            text_content = data.decode(errors="ignore")

        elif ext == "dcm":
            dcm = pydicom.dcmread(io.BytesIO(data))
            text_content = str(dcm)  # Добавляем метаданные как текст
            if hasattr(dcm, "pixel_array"):
                img = Image.fromarray(dcm.pixel_array).convert("L")
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                images_base64.append(base64.b64encode(buf.getvalue()).decode('utf-8'))

        elif ext in ("png", "jpg", "jpeg", "tiff", "bmp", "gif"):
            images_base64.append(base64.b64encode(data).decode('utf-8'))

    except Exception as e:
        st.warning(f"Не удалось полностью обработать файл {fname}: {e}")

    return {"text": text_content, "images": images_base64}


def integrated_analysis() -> str:
    """Собираем все файлы, заметки, изображения и отправляем запрос к модели."""
    header = (
        f"Пациент: {patient_name or '—'}, {patient_age or '—'} лет, пол: {patient_sex or '—'}.\n"
        "Вы — ведущий врач‑консультант с 20‑летним стажем, специалист по комплексной диагностике. Ваша задача — провести глубокий и всесторонний анализ предоставленных медицинских данных.\n\n"
        "ТРЕБОВАНИЯ К АНАЛИЗУ:\n"
        "1.  **Анализ в динамике:** Обязательно сравнивайте результаты анализов, сделанных в разные даты. Отмечайте любые тенденции (улучшение, ухудшение, стабильность).\n"
        "2.  **Анализ изображений:** Внимательно изучите все предоставленные изображения (графики, снимки, диаграммы). Опишите, что вы на них видите, и как это соотносится с текстовыми данными.\n"
        "3.  **Глубина и детализация:** Ответ должен быть максимально подробным. Объясняйте свои выводы, ссылаясь на конкретные показатели в анализах.\n\n"
        "СФОРМИРУЙТЕ ЗАКЛЮЧЕНИЕ СТРОГО ПО СТРУКТУРЕ:\n"
        "1) Проведённые исследования и анализы (перечислите ключевые исследования, которые вы проанализировали).\n"
        "2) Заключение по результатам (детальный разбор с указанием динамики и интерпретацией визуальных данных).\n"
        "3) Предварительный диагноз или гипотеза (если данных достаточно).\n"
        "4) План лечения (конкретные препараты, дозировки, кратность приема и длительность курса).\n"
        "5) Рекомендации (изменение образа жизни, дополнительные обследования с указанием цели, рекомендуемые сроки для контрольных анализов).\n\n"
        "ЗАПРЕЩЕНО: Использовать списки с маркерами (•, -, *), эмодзи, заголовки в ВЕРХНЕМ РЕГИСТРЕ, фразы «обратитесь к врачу» или упоминания, что вы ИИ."
    )

    messages = [{"role": "system", "content": header}]
    user_content = []

    if global_note.strip():
        user_content.append({"type": "text", "text": f"Общие указания от лечащего врача: {global_note.strip()}"})

    for fname, meta in st.session_state.files.items():
        file_date = meta['date'].strftime('%d.%m.%Y')
        file_note = meta['note'].strip()

        file_header = f"=== Анализ файла: {fname} (дата: {file_date}) ===\n"
        if file_note:
            file_header += f"Комментарий к файлу: {file_note}\n"

        processed_file = process_file(fname, meta["data"])
        file_text = processed_file['text'].strip()

        if file_text:
            user_content.append({"type": "text", "text": file_header + file_text[:20000]})
        else:
            user_content.append(
                {"type": "text", "text": file_header + "(Файл не содержит текста, см. изображение ниже)"})

        for img_b64 in processed_file['images']:
            user_content.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{img_b64}"}
            })

    if not user_content:
        user_content.append({"type": "text", "text": "Нет данных для анализа."})

    messages.append({"role": "user", "content": user_content})

    resp = openai_client.chat.completions.create(
        model=MODEL,
        messages=messages,
        temperature=TEMPERATURE,
        max_tokens=MAX_TOKENS
    )
    return resp.choices[0].message.content.strip()


def build_docx(text: str) -> bytes:
    doc = docx.Document()
    p = doc.add_paragraph()
    r = p.add_run("Медицинское заключение")
    r.bold = True;
    r.font.size = Pt(22);
    r.font.color.rgb = THEME
    doc.add_paragraph(f"{datetime.datetime.now():%d.%m.%Y %H:%M}")
    doc.add_paragraph(f"Пациент: {patient_name or '—'}, {patient_age or '—'} лет, пол: {patient_sex or '—'}")
    doc.add_paragraph().add_run("═" * 40)

    for line in text.splitlines():
        doc.add_paragraph(line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, f"{uuid.uuid4()}.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        try:
            convert(docx_path, tmp)
            pdf_paths = glob.glob(os.path.join(tmp, "*.pdf"))
            if not pdf_paths:
                raise FileNotFoundError("Конвертация в PDF не удалась, файл не был создан.")
            return open(pdf_paths[0], "rb").read()
        except Exception as e:
            st.error(f"Ошибка конвертации в PDF: {e}. Убедитесь, что на сервере установлен MS Word или LibreOffice.")
            st.stop()


# ---------- Генерация отчёта ----------
def generate_report(to_pdf: bool):
    if not st.session_state.files:
        st.warning("⚠️ Загрузите хотя бы один файл для анализа.")
        return
    with st.spinner("ИИ проводит комплексный анализ материалов... Это может занять несколько минут..."):
        result_text = integrated_analysis()

    st.success("Анализ завершен! Отчёт готов к скачиванию.")
    st.markdown("### Сгенерированный отчёт:")
    st.text_area("Текст отчёта", result_text, height=600)

    docx_bytes = build_docx(result_text)

    col1, col2, _ = st.columns([1, 1, 3])

    with col1:
        st.download_button(
            "⬇️ Скачать DOCX",
            docx_bytes,
            f"MedConsultant_Report_{patient_name or 'P'}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="docx_download"
        )

    if to_pdf:
        with st.spinner("Конвертируем в PDF..."):
            pdf_bytes = docx_to_pdf(docx_bytes)
            with col2:
                st.download_button(
                    "⬇️ Скачать PDF",
                    pdf_bytes,
                    f"MedConsultant_Report_{patient_name or 'P'}.pdf",
                    mime="application/pdf",
                    key="pdf_download"
                )


# ---------- Кнопки управления ----------
st.markdown("---")
st.subheader("🚀 Сгенерировать отчёт")
col1, col2 = st.columns(2)
with col1:
    if st.button("📄 Создать и скачать DOCX‑отчёт", type="primary"):
        generate_report(False)
with col2:
    if st.button("📑 Создать и скачать PDF‑отчёт"):
        generate_report(True)